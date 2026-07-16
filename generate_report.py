"""
DevFlow – Complete Professional Project Report Generator (.docx)
================================================================
Generates a comprehensive, university-grade project report covering
every aspect of the DevFlow project management system.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime
import os

# ── Colour Palette ──────────────────────────────────────────────
NAVY      = RGBColor(0x1E, 0x3A, 0x8A)
TEAL      = RGBColor(0x0D, 0x94, 0x88)
DARK_GRAY = RGBColor(0x33, 0x41, 0x55)
LIGHT_BG  = "F8FAFC"
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)

# ── Helper Functions ────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_heading_styled(doc, text, level=1):
    """Add a heading with custom styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level <= 2 else TEAL
    return h


def add_body(doc, text):
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.4
    return p


def add_code_block(doc, code_text, caption=""):
    """Add a code block in a shaded table cell."""
    if caption:
        cp = doc.add_paragraph()
        run = cp.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = TEAL

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BG)
    p = cell.paragraphs[0]
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = DARK_GRAY
    # set font
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Consolas" w:hAnsi="Consolas"/>')
    rPr.append(rFonts)
    doc.add_paragraph()  # spacer


def add_bullet(doc, text, bold_prefix=""):
    """Add a bullet-point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def make_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1E3A8A")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            if r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_BG)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table


def read_file(path):
    """Read a source file and return its contents."""
    full_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), path)
    if not os.path.exists(full_path):
        return f"# File not found: {path}"
    with open(full_path, "r", encoding="utf-8") as f:
        return f.read()


# ── Main Report Builder ─────────────────────────────────────────

def build_report():
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    today = datetime.datetime.now().strftime("%B %d, %Y")

    # ══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DEVFLOW")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Agile Project Management System")
    run.font.size = Pt(22)
    run.font.color.rgb = TEAL

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Comprehensive Project Report")
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_GRAY

    for _ in range(4):
        doc.add_paragraph()

    cover_info = [
        ("Project Title:", "DevFlow – Agile Project Management System"),
        ("Technology Stack:", "Django REST Framework · Next.js 16 · MongoDB Atlas · Docker"),
        ("Architecture:", "Full-Stack Web Application with REST API & Role-Based Access Control"),
        ("Design Patterns:", "Singleton, Factory, Builder, Strategy, State, Observer, Command, Chain of Responsibility, Facade"),
        ("Report Date:", today),
        ("Document Type:", "Professional Technical Project Report"),
    ]
    for label, value in cover_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        run = p.add_run(value)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1.  Introduction",
        "2.  Project Objectives",
        "3.  Scope of the Project",
        "4.  System Architecture Overview",
        "5.  Technology Stack",
        "6.  Project Directory Structure",
        "7.  Database Design (MongoDB Atlas)",
        "    7.1  User Collection",
        "    7.2  Project Collection",
        "    7.3  Sprint Collection",
        "    7.4  Task Collection",
        "    7.5  Bug Collection",
        "    7.6  Notification Collection",
        "    7.7  Entity-Relationship Summary",
        "8.  Backend Architecture & Modules",
        "    8.1  Accounts Module",
        "    8.2  Projects Module",
        "    8.3  Sprints Module",
        "    8.4  Tasks Module",
        "    8.5  Bugs Module",
        "    8.6  Notifications Module",
        "    8.7  Reports & Dashboard Module",
        "    8.8  Core Module",
        "9.  Design Patterns Implementation",
        "    9.1  Singleton Pattern",
        "    9.2  Factory Pattern",
        "    9.3  Builder Pattern",
        "    9.4  Strategy Pattern",
        "    9.5  State Pattern",
        "    9.6  Observer Pattern",
        "    9.7  Command Pattern",
        "    9.8  Chain of Responsibility Pattern",
        "    9.9  Facade Pattern",
        "10. REST API Documentation",
        "11. Authentication & Authorization (JWT)",
        "12. Role-Based Access Control (RBAC)",
        "13. Frontend Architecture (Next.js)",
        "    13.1 Pages & Routing",
        "    13.2 Components",
        "    13.3 Services Layer",
        "    13.4 Middleware & Auth Guard",
        "    13.5 UI/UX Design System",
        "14. Containerization & Deployment (Docker)",
        "15. Environment Configuration",
        "16. Testing Strategy",
        "17. Conclusion",
        "18. Future Enhancements",
        "19. References",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 1 – INTRODUCTION
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Introduction", level=1)
    add_body(doc,
        "DevFlow is a comprehensive, full-stack Agile Project Management System designed to facilitate "
        "efficient software development lifecycle management. The system provides a centralized platform "
        "for teams to manage projects, plan sprints, track tasks through Kanban-style workflows, "
        "report and monitor bugs, and generate insightful project reports — all within a secure, "
        "role-based access environment."
    )
    add_body(doc,
        "The application follows a modern client-server architecture with a Django REST Framework "
        "backend connected to MongoDB Atlas (a cloud-hosted NoSQL database), and a Next.js 16 frontend "
        "powered by React 19 and TailwindCSS v4. The system implements nine (9) industry-standard "
        "software design patterns to ensure maintainability, extensibility, and adherence to "
        "SOLID principles of object-oriented design."
    )
    add_body(doc,
        "DevFlow is containerized using Docker and Docker Compose, enabling seamless deployment "
        "across development, staging, and production environments. The system supports five distinct "
        "user roles — Admin, Project Manager, Developer, QA Engineer, and Client — each with "
        "specific permissions and access levels enforced at both the API and UI layers."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 2 – PROJECT OBJECTIVES
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. Project Objectives", level=1)
    add_body(doc, "The primary objectives of the DevFlow project are as follows:")
    objectives = [
        "To develop a full-stack web application that supports the complete Agile project management lifecycle including project creation, sprint planning, task tracking, and bug management.",
        "To implement Role-Based Access Control (RBAC) with five distinct user roles (Admin, Project Manager, Developer, QA, Client) ensuring secure and appropriate access to system features.",
        "To integrate MongoDB Atlas as the primary database using MongoEngine ODM for flexible, schema-less document storage suitable for agile workflows.",
        "To implement nine (9) software design patterns (Singleton, Factory, Builder, Strategy, State, Observer, Command, Chain of Responsibility, and Facade) demonstrating enterprise-grade software architecture.",
        "To build a RESTful API backend using Django REST Framework with JWT-based authentication for secure, stateless communication between client and server.",
        "To develop a modern, responsive, and visually appealing frontend using Next.js 16 with React 19, featuring glassmorphism design, dark theme, and micro-animations.",
        "To containerize the entire application using Docker and Docker Compose for consistent deployment and environment isolation.",
        "To implement a Kanban-style task state machine with defined transitions (Todo → InProgress → Testing → Done) governed by the State Design Pattern.",
        "To provide real-time dashboard analytics and report generation using the Builder and Facade patterns for aggregated project metrics.",
        "To implement a notification system using the Observer pattern for event-driven communication across the application.",
    ]
    for obj in objectives:
        add_bullet(doc, obj)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 3 – SCOPE
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Scope of the Project", level=1)

    add_heading_styled(doc, "3.1 In-Scope Features", level=2)
    in_scope = [
        "User Registration and Login with JWT Authentication (Access + Refresh Tokens)",
        "Role-based user management (Admin, ProjectManager, Developer, QA, Client)",
        "Full CRUD operations for Projects, Sprints, Tasks, and Bugs",
        "Task status workflow management with State Pattern (Todo → InProgress → Testing → Done)",
        "Task priority management using Strategy Pattern (Low, Medium, High)",
        "Bug lifecycle management (Open → Assigned → Resolved → Verified → Closed)",
        "Sprint lifecycle management (Planned → Active → Completed)",
        "Dashboard with aggregated statistics (project, sprint, task, bug counts)",
        "Report generation using Builder Pattern",
        "Notification system with Observer Pattern",
        "Approval chain using Chain of Responsibility Pattern",
        "Email notification service integration",
        "Frontend with responsive Glassmorphism UI design",
        "Protected routes with middleware authentication guards",
        "Docker containerization for backend and MongoDB services",
    ]
    for item in in_scope:
        add_bullet(doc, item)

    add_heading_styled(doc, "3.2 Out-of-Scope (Future Considerations)", level=2)
    out_scope = [
        "Real-time WebSocket notifications (currently print-based observers)",
        "File attachment and media upload for tasks/bugs",
        "Advanced analytics with charts driven by real time-series data",
        "Multi-tenant / multi-organization support",
        "Mobile application (iOS/Android)",
        "CI/CD pipeline automation",
    ]
    for item in out_scope:
        add_bullet(doc, item)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 4 – SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. System Architecture Overview", level=1)
    add_body(doc,
        "DevFlow follows a three-tier client-server architecture consisting of a Presentation Layer (Frontend), "
        "an Application Logic Layer (Backend API), and a Data Layer (MongoDB Atlas). The architecture is designed "
        "for separation of concerns, scalability, and maintainability."
    )

    add_heading_styled(doc, "4.1 Architecture Diagram (Textual)", level=2)
    arch_text = """
┌──────────────────────────────────────────────────────────────────────┐
│                      PRESENTATION LAYER                             │
│  ┌───────────────────────────────────────────────────────────────┐  │
│  │  Next.js 16 Frontend (React 19 + TypeScript + TailwindCSS v4)│  │
│  │  ├── Pages: Login, Register, Dashboard, Projects, Tasks, Bugs│  │
│  │  ├── Components: Navbar, Sidebar, Icons, Cards               │  │
│  │  ├── Services: auth.ts, project.ts, task.ts, bug.ts, api.ts  │  │
│  │  └── Middleware: JWT Auth Guard (middleware.ts)               │  │
│  └──────────────────────────┬────────────────────────────────────┘  │
│                             │ HTTP REST (JSON)                      │
│                             ▼                                       │
│                   APPLICATION LOGIC LAYER                           │
│  ┌───────────────────────────────────────────────────────────────┐  │
│  │  Django REST Framework Backend (Python 3.12)                  │  │
│  │  ├── accounts/ (Register, Login, JWT, Permissions)            │  │
│  │  ├── projects/ (CRUD, Service Layer)                          │  │
│  │  ├── sprints/ (CRUD, Close Sprint, Service Layer)             │  │
│  │  ├── tasks/ (CRUD, State Machine, Strategy, Service Layer)    │  │
│  │  ├── bugs/ (CRUD, Assign, Resolve, Verify, Service Layer)    │  │
│  │  ├── notifications/ (Observer, Subject, Email Service)        │  │
│  │  ├── reports/ (Builder, Dashboard Facade, Service Layer)      │  │
│  │  └── core/ (Singleton, Factory, Command, Chain, Facade, RBAC)│  │
│  └──────────────────────────┬────────────────────────────────────┘  │
│                             │ MongoEngine ODM                       │
│                             ▼                                       │
│                        DATA LAYER                                   │
│  ┌───────────────────────────────────────────────────────────────┐  │
│  │  MongoDB Atlas (Cloud Cluster)                                │  │
│  │  Database: devflow                                            │  │
│  │  Collections: user, projects, sprints, tasks, bugs,           │  │
│  │               notification                                    │  │
│  └───────────────────────────────────────────────────────────────┘  │
└──────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, arch_text.strip(), "System Architecture:")

    add_heading_styled(doc, "4.2 Communication Flow", level=2)
    add_body(doc,
        "1. The user interacts with the Next.js frontend, which runs on port 3000.\n"
        "2. The frontend sends HTTP REST API requests (JSON) to the Django backend on port 8000.\n"
        "3. JWT tokens (access + refresh) are sent via the Authorization header (Bearer token).\n"
        "4. The backend authenticates requests using a custom MongoJWTAuthentication class.\n"
        "5. The backend processes business logic through Service classes and Design Patterns.\n"
        "6. Data is persisted to MongoDB Atlas via MongoEngine ODM (Document models).\n"
        "7. Responses are returned as JSON to the frontend for rendering."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 5 – TECHNOLOGY STACK
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Technology Stack", level=1)

    add_heading_styled(doc, "5.1 Backend Technologies", level=2)
    make_table(doc,
        ["Technology", "Version", "Purpose"],
        [
            ["Python", "3.12", "Core programming language for backend development"],
            ["Django", "Latest", "High-level Python web framework providing URL routing, ORM, and middleware"],
            ["Django REST Framework", "Latest", "Toolkit for building Web APIs with serializers, views, and authentication"],
            ["MongoEngine", "Latest", "MongoDB Object-Document Mapper (ODM) for Python"],
            ["PyMongo", "Latest", "Low-level MongoDB driver used internally by MongoEngine"],
            ["SimpleJWT", "Latest", "JWT token generation and validation for DRF authentication"],
            ["bcrypt", "Latest", "Industry-standard password hashing library"],
            ["django-cors-headers", "Latest", "Handles Cross-Origin Resource Sharing for frontend-backend communication"],
            ["python-dotenv", "Latest", "Loads environment variables from .env files"],
        ],
        col_widths=[4, 2.5, 10]
    )

    add_heading_styled(doc, "5.2 Frontend Technologies", level=2)
    make_table(doc,
        ["Technology", "Version", "Purpose"],
        [
            ["Next.js", "16.2.9", "React-based framework with App Router, SSR, and file-based routing"],
            ["React", "19.2.4", "Component-based UI library for building interactive interfaces"],
            ["TypeScript", "5.x", "Statically typed superset of JavaScript for type-safe development"],
            ["TailwindCSS", "4.3.0", "Utility-first CSS framework for rapid UI development"],
            ["Axios", "1.17.0", "HTTP client library (available but fetch API primarily used)"],
        ],
        col_widths=[4, 2.5, 10]
    )

    add_heading_styled(doc, "5.3 Database & Infrastructure", level=2)
    make_table(doc,
        ["Technology", "Purpose"],
        [
            ["MongoDB Atlas", "Cloud-hosted NoSQL document database (Replica Set cluster)"],
            ["SQLite3", "Django's default database for admin/session tables (internal use)"],
            ["Docker", "Container platform for application packaging and isolation"],
            ["Docker Compose", "Multi-container orchestration (backend + MongoDB services)"],
        ],
        col_widths=[5, 11.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 6 – DIRECTORY STRUCTURE
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Project Directory Structure", level=1)
    add_body(doc, "The DevFlow project is organized into two main directories: backend (Django) and frontend (Next.js), along with Docker configuration files at the root level.")

    dir_tree = """
DevFlow/
├── .env                         # Environment variables (MONGO_URI, SECRET_KEY, DEBUG)
├── Dockerfile                   # Docker image definition (Python 3.12)
├── docker-compose.yml           # Multi-service orchestration (backend + MongoDB)
├── requirements.txt             # Python dependencies
│
├── backend/
│   ├── manage.py                # Django management entry point
│   ├── db.sqlite3               # SQLite DB for Django admin internals
│   ├── generate_pdf.py          # Design patterns PDF report generator
│   │
│   ├── devflow/                 # Django project configuration
│   │   ├── __init__.py
│   │   ├── settings.py          # Settings: INSTALLED_APPS, REST_FRAMEWORK, CORS, DB
│   │   ├── urls.py              # Root URL configuration (all API routes)
│   │   ├── asgi.py              # ASGI entry point
│   │   └── wsgi.py              # WSGI entry point
│   │
│   ├── accounts/                # User authentication & management module
│   │   ├── models.py            # User Document model (MongoEngine)
│   │   ├── views.py             # RegisterView, LoginView (API endpoints)
│   │   ├── serializers.py       # RegisterSerializer (DRF validation)
│   │   ├── services.py          # PasswordService (bcrypt hashing)
│   │   ├── jwt_service.py       # JWTService (token generation)
│   │   ├── authentication.py    # MongoJWTAuthentication (custom DRF auth backend)
│   │   ├── permissions.py       # Role-based permission classes (5 roles)
│   │   ├── admin_views.py       # Admin-only UsersListView
│   │   └── urls.py              # /register/, /login/
│   │
│   ├── projects/                # Project management module
│   │   ├── models.py            # Project Document model
│   │   ├── views.py             # CRUD views (Create, Get, GetAll, Update, Delete)
│   │   ├── serializers.py       # ProjectSerializer
│   │   ├── services.py          # ProjectService (business logic)
│   │   └── urls.py              # Project API routes
│   │
│   ├── sprints/                 # Sprint management module
│   │   ├── models.py            # Sprint Document model
│   │   ├── views.py             # CRUD + CloseSprint views
│   │   ├── serializers.py       # SprintSerializer
│   │   ├── services.py          # SprintService (business logic)
│   │   └── urls.py              # Sprint API routes
│   │
│   ├── tasks/                   # Task management module
│   │   ├── models.py            # Task Document model (priority + status enums)
│   │   ├── views.py             # CRUD + ChangeTaskStatus views
│   │   ├── serializers.py       # TaskSerializer
│   │   ├── services.py          # TaskService (business logic)
│   │   ├── state.py             # State Pattern (TodoState → DoneState)
│   │   ├── strategy.py          # Strategy Pattern (Priority strategies)
│   │   └── urls.py              # Task API routes
│   │
│   ├── bugs/                    # Bug tracking module
│   │   ├── models.py            # Bug Document model (severity + status enums)
│   │   ├── views.py             # CRUD + Assign + Resolve + Verify views
│   │   ├── serializers.py       # BugSerializer
│   │   ├── services.py          # BugService (business logic)
│   │   └── urls.py              # Bug API routes
│   │
│   ├── notifications/           # Notification module
│   │   ├── models.py            # Notification Document model
│   │   ├── views.py             # NotificationListView, MarkAsReadView
│   │   ├── observer.py          # Observer Pattern (NotificationObserver)
│   │   ├── subject.py           # Subject class (NotificationSubject)
│   │   ├── email_service.py     # EmailService (Django send_mail)
│   │   └── urls.py              # Notification API routes
│   │
│   ├── reports/                 # Reports & Dashboard module
│   │   ├── builder.py           # Builder Pattern (ReportBuilder + Report)
│   │   ├── dashboard.py         # DashboardView (uses Facade)
│   │   ├── services.py          # ReportService (uses Builder)
│   │   ├── views.py             # GenerateReportView
│   │   └── urls.py              # /generate/, /dashboard/
│   │
│   └── core/                    # Shared core utilities & design patterns
│       ├── db.py                # MongoDB connection (MongoEngine connect)
│       ├── singleton.py         # Singleton Pattern (MongoConnection)
│       ├── factories.py         # Factory Pattern (UserFactory)
│       ├── commands.py          # Command Pattern (CreateTask, DeleteTask)
│       ├── approval_chain.py    # Chain of Responsibility (PM → Admin)
│       ├── observers.py         # Observer Pattern (Email + Notification observers)
│       ├── services.py          # Facade Pattern (DashboardFacade)
│       └── permissions.py       # Core RBAC permission classes
│
└── frontend/
    ├── package.json             # NPM dependencies and scripts
    ├── next.config.ts           # Next.js configuration
    ├── tsconfig.json            # TypeScript configuration
    ├── middleware.ts             # Auth guard middleware (JWT cookie check)
    ├── eslint.config.mjs        # ESLint configuration
    ├── postcss.config.mjs       # PostCSS configuration (TailwindCSS)
    │
    ├── app/                     # Next.js App Router pages
    │   ├── layout.tsx           # Root layout (Geist fonts, global CSS)
    │   ├── page.tsx             # Home page (redirects to /dashboard)
    │   ├── globals.css          # Global styles (Glassmorphism, dark theme, animations)
    │   ├── login/page.tsx       # Login page
    │   ├── register/page.tsx    # Registration page (with role selection)
    │   ├── dashboard/page.tsx   # Dashboard (stats, charts, quick actions)
    │   ├── projects/page.tsx    # Projects management page
    │   ├── tasks/page.tsx       # Tasks management page (Kanban)
    │   └── bugs/page.tsx        # Bugs tracking page
    │
    ├── components/              # Reusable UI components
    │   ├── Navbar.tsx           # Top navigation bar (logo, logout)
    │   ├── Sidebar.tsx          # Left sidebar navigation (menu, user info)
    │   ├── Icons.tsx            # SVG icon components (11 icons)
    │   ├── ProjectCard.tsx      # Project display card
    │   ├── TaskCard.tsx         # Task display card
    │   └── BugCard.tsx          # Bug display card
    │
    ├── services/                # API service functions
    │   ├── api.ts               # Base API URL configuration
    │   ├── auth.ts              # Login, register, getAuthHeaders
    │   ├── project.ts           # CRUD operations for projects
    │   ├── task.ts              # CRUD + status change for tasks
    │   └── bug.ts               # CRUD + assign + resolve + verify for bugs
    │
    └── public/                  # Static assets
"""
    add_code_block(doc, dir_tree.strip(), "Complete Directory Tree:")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 7 – DATABASE DESIGN
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. Database Design (MongoDB Atlas)", level=1)
    add_body(doc,
        "DevFlow uses MongoDB Atlas, a cloud-hosted NoSQL document database, as its primary data store. "
        "The database is named 'devflow' and is hosted on a MongoDB Atlas replica set cluster with TLS "
        "encryption and authentication. The application uses MongoEngine as the ODM (Object-Document Mapper) "
        "to define and interact with document schemas in Python."
    )
    add_body(doc,
        "MongoDB was chosen for its schema flexibility, which aligns well with Agile development workflows "
        "where data structures may evolve over time. Each collection stores documents that correspond to "
        "a specific domain entity."
    )

    # 7.1 User
    add_heading_styled(doc, "7.1 User Collection", level=2)
    make_table(doc,
        ["Field Name", "Data Type", "Constraints", "Description"],
        [
            ["_id", "ObjectId", "Primary Key (auto)", "Unique document identifier"],
            ["name", "StringField", "Required", "Full name of the user"],
            ["email", "EmailField", "Required, Unique", "Email address (used for login)"],
            ["password", "StringField", "Required", "Bcrypt-hashed password"],
            ["role", "StringField", "Choices: Admin, ProjectManager, Developer, QA, Client", "Role determining access permissions"],
            ["created_at", "DateTimeField", "Default: datetime.utcnow", "Account creation timestamp"],
        ],
        col_widths=[3, 3, 5, 5.5]
    )
    add_code_block(doc, read_file("backend/accounts/models.py"), "Source Code — accounts/models.py:")

    # 7.2 Project
    add_heading_styled(doc, "7.2 Project Collection", level=2)
    make_table(doc,
        ["Field Name", "Data Type", "Constraints", "Description"],
        [
            ["_id", "ObjectId", "Primary Key (auto)", "Unique document identifier"],
            ["name", "StringField", "Required", "Project name"],
            ["description", "StringField", "Optional", "Detailed project description"],
            ["status", "StringField", "Optional", "Current project status"],
            ["manager_id", "StringField", "Optional", "Reference to the assigned Project Manager (User ID)"],
        ],
        col_widths=[3, 3, 4, 6.5]
    )
    add_code_block(doc, read_file("backend/projects/models.py"), "Source Code — projects/models.py:")

    # 7.3 Sprint
    add_heading_styled(doc, "7.3 Sprint Collection", level=2)
    make_table(doc,
        ["Field Name", "Data Type", "Constraints", "Description"],
        [
            ["_id", "ObjectId", "Primary Key (auto)", "Unique document identifier"],
            ["project_id", "StringField", "Required", "Reference to parent Project ID"],
            ["name", "StringField", "Required", "Sprint name/identifier"],
            ["goal", "StringField", "Optional", "Sprint goal/objective"],
            ["start_date", "DateTimeField", "Optional", "Sprint start date"],
            ["end_date", "DateTimeField", "Optional", "Sprint end date"],
            ["status", "StringField", "Default: 'Planned'", "Sprint status (Planned/Active/Completed)"],
            ["created_at", "DateTimeField", "Default: utcnow", "Sprint creation timestamp"],
        ],
        col_widths=[3, 3, 4, 6.5]
    )
    add_code_block(doc, read_file("backend/sprints/models.py"), "Source Code — sprints/models.py:")

    # 7.4 Task
    add_heading_styled(doc, "7.4 Task Collection", level=2)
    make_table(doc,
        ["Field Name", "Data Type", "Constraints", "Description"],
        [
            ["_id", "ObjectId", "Primary Key (auto)", "Unique document identifier"],
            ["sprint_id", "StringField", "Required", "Reference to parent Sprint ID"],
            ["title", "StringField", "Required", "Task title/name"],
            ["description", "StringField", "Optional", "Detailed task description"],
            ["assigned_to", "StringField", "Optional", "Developer User ID assigned to task"],
            ["priority", "StringField", "Choices: Low/Medium/High, Default: Medium", "Task priority level"],
            ["status", "StringField", "Choices: Todo/InProgress/Testing/Done, Default: Todo", "Current task status (State Pattern)"],
            ["created_at", "DateTimeField", "Default: utcnow", "Task creation timestamp"],
        ],
        col_widths=[3, 3, 5, 5.5]
    )
    add_code_block(doc, read_file("backend/tasks/models.py"), "Source Code — tasks/models.py:")

    # 7.5 Bug
    add_heading_styled(doc, "7.5 Bug Collection", level=2)
    make_table(doc,
        ["Field Name", "Data Type", "Constraints", "Description"],
        [
            ["_id", "ObjectId", "Primary Key (auto)", "Unique document identifier"],
            ["task_id", "StringField", "Required", "Reference to related Task ID"],
            ["title", "StringField", "Required", "Bug title/summary"],
            ["description", "StringField", "Optional", "Detailed bug description"],
            ["reported_by", "StringField", "Optional", "User ID who reported the bug"],
            ["assigned_to", "StringField", "Optional", "Developer User ID assigned to fix"],
            ["severity", "StringField", "Choices: Low/Medium/High/Critical, Default: Low", "Bug severity level"],
            ["status", "StringField", "Choices: Open/Assigned/Resolved/Verified/Closed, Default: Open", "Current bug lifecycle status"],
            ["created_at", "DateTimeField", "Default: utcnow", "Bug report timestamp"],
        ],
        col_widths=[3, 3, 5, 5.5]
    )
    add_code_block(doc, read_file("backend/bugs/models.py"), "Source Code — bugs/models.py:")

    # 7.6 Notification
    add_heading_styled(doc, "7.6 Notification Collection", level=2)
    make_table(doc,
        ["Field Name", "Data Type", "Constraints", "Description"],
        [
            ["_id", "ObjectId", "Primary Key (auto)", "Unique document identifier"],
            ["user_id", "StringField", "Optional", "Target user for the notification"],
            ["title", "StringField", "Optional", "Notification title"],
            ["message", "StringField", "Optional", "Notification message content"],
            ["is_read", "BooleanField", "Default: False", "Read/unread status flag"],
            ["created_at", "DateTimeField", "Default: utcnow", "Notification creation timestamp"],
        ],
        col_widths=[3, 3, 4, 6.5]
    )
    add_code_block(doc, read_file("backend/notifications/models.py"), "Source Code — notifications/models.py:")

    # 7.7 ER Summary
    add_heading_styled(doc, "7.7 Entity-Relationship Summary", level=2)
    add_body(doc, "The following table summarizes the relationships between entities in the DevFlow system:")
    make_table(doc,
        ["Relationship", "Type", "Description"],
        [
            ["User → Project", "One-to-Many", "A Project Manager (User) can manage multiple projects (via manager_id)"],
            ["Project → Sprint", "One-to-Many", "Each project can have multiple sprints (via project_id)"],
            ["Sprint → Task", "One-to-Many", "Each sprint contains multiple tasks (via sprint_id)"],
            ["Task → Bug", "One-to-Many", "Each task can have multiple associated bugs (via task_id)"],
            ["User → Task", "One-to-Many", "A Developer (User) can be assigned multiple tasks (via assigned_to)"],
            ["User → Bug", "Many-to-Many", "Users can report bugs (reported_by) and be assigned to fix them (assigned_to)"],
            ["User → Notification", "One-to-Many", "Each user can have multiple notifications (via user_id)"],
        ],
        col_widths=[4, 3, 9.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 8 – BACKEND ARCHITECTURE
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "8. Backend Architecture & Modules", level=1)
    add_body(doc,
        "The DevFlow backend follows a modular architecture with each Django app responsible for a specific "
        "domain. Each module follows a consistent pattern: Models → Serializers → Services → Views → URLs. "
        "This layered approach separates data modeling, validation, business logic, API handling, and routing."
    )

    # 8.1 Accounts
    add_heading_styled(doc, "8.1 Accounts Module", level=2)
    add_body(doc,
        "The Accounts module handles all user-related operations including registration, login, JWT token "
        "management, password hashing, and role-based permission enforcement. It is the security backbone of the application."
    )
    add_body(doc, "Key Components:")
    add_bullet(doc, "User model (MongoEngine Document) with name, email, password, role, and created_at fields.", "models.py — ")
    add_bullet(doc, "RegisterView handles new user creation with bcrypt password hashing.", "views.py — ")
    add_bullet(doc, "LoginView authenticates users and returns JWT access + refresh tokens.", "views.py — ")
    add_bullet(doc, "MongoJWTAuthentication is a custom DRF authentication backend that validates JWT tokens and resolves MongoDB users.", "authentication.py — ")
    add_bullet(doc, "JWTService generates JWT tokens with user_id and role claims embedded in the payload.", "jwt_service.py — ")
    add_bullet(doc, "PasswordService provides static methods for bcrypt hashing and verification.", "services.py — ")
    add_bullet(doc, "Five permission classes (IsAdmin, IsProjectManager, IsDeveloper, IsQA, IsClient) for RBAC.", "permissions.py — ")
    add_bullet(doc, "Admin-only endpoint to list all users in the system.", "admin_views.py — ")

    add_code_block(doc, read_file("backend/accounts/authentication.py"), "Source Code — accounts/authentication.py (Custom JWT Auth Backend):")
    add_code_block(doc, read_file("backend/accounts/jwt_service.py"), "Source Code — accounts/jwt_service.py:")
    add_code_block(doc, read_file("backend/accounts/views.py"), "Source Code — accounts/views.py:")

    # 8.2 Projects
    add_heading_styled(doc, "8.2 Projects Module", level=2)
    add_body(doc,
        "The Projects module provides complete CRUD operations for project management. Only users with "
        "the ProjectManager role can create projects (enforced by IsProjectManager permission class). "
        "All project data is stored in the 'projects' MongoDB collection."
    )
    add_code_block(doc, read_file("backend/projects/views.py"), "Source Code — projects/views.py:")
    add_code_block(doc, read_file("backend/projects/services.py"), "Source Code — projects/services.py:")

    # 8.3 Sprints
    add_heading_styled(doc, "8.3 Sprints Module", level=2)
    add_body(doc,
        "The Sprints module manages sprint lifecycle within projects. It supports CRUD operations plus "
        "a dedicated CloseSprintView that sets a sprint's status to 'Completed'. Sprints link to projects "
        "via project_id and contain a goal, start/end dates, and status."
    )
    add_code_block(doc, read_file("backend/sprints/views.py"), "Source Code — sprints/views.py:")
    add_code_block(doc, read_file("backend/sprints/services.py"), "Source Code — sprints/services.py:")

    # 8.4 Tasks
    add_heading_styled(doc, "8.4 Tasks Module", level=2)
    add_body(doc,
        "The Tasks module is the most feature-rich module in the backend. It manages task CRUD operations, "
        "implements the State Pattern for task status transitions (Todo → InProgress → Testing → Done), "
        "and provides the Strategy Pattern for priority handling. Only Developers can change task status "
        "(enforced by IsDeveloper permission)."
    )
    add_code_block(doc, read_file("backend/tasks/views.py"), "Source Code — tasks/views.py:")
    add_code_block(doc, read_file("backend/tasks/state.py"), "Source Code — tasks/state.py (State Pattern):")
    add_code_block(doc, read_file("backend/tasks/strategy.py"), "Source Code — tasks/strategy.py (Strategy Pattern):")

    # 8.5 Bugs
    add_heading_styled(doc, "8.5 Bugs Module", level=2)
    add_body(doc,
        "The Bugs module provides comprehensive bug lifecycle management. Bugs are linked to tasks via "
        "task_id and progress through a defined lifecycle: Open → Assigned → Resolved → Verified → Closed. "
        "Bug verification is restricted to QA role users (enforced by IsQA permission)."
    )
    add_code_block(doc, read_file("backend/bugs/views.py"), "Source Code — bugs/views.py:")
    add_code_block(doc, read_file("backend/bugs/services.py"), "Source Code — bugs/services.py:")

    # 8.6 Notifications
    add_heading_styled(doc, "8.6 Notifications Module", level=2)
    add_body(doc,
        "The Notifications module implements an event-driven notification system using the Observer Pattern. "
        "It provides API endpoints to list all notifications and mark individual notifications as read. "
        "The module includes NotificationSubject (subject) and NotificationObserver (observer) classes "
        "for the Observer pattern, and an EmailService for email notifications."
    )
    add_code_block(doc, read_file("backend/notifications/views.py"), "Source Code — notifications/views.py:")
    add_code_block(doc, read_file("backend/notifications/observer.py"), "Source Code — notifications/observer.py:")
    add_code_block(doc, read_file("backend/notifications/subject.py"), "Source Code — notifications/subject.py:")
    add_code_block(doc, read_file("backend/notifications/email_service.py"), "Source Code — notifications/email_service.py:")

    # 8.7 Reports
    add_heading_styled(doc, "8.7 Reports & Dashboard Module", level=2)
    add_body(doc,
        "The Reports module provides dashboard analytics and report generation capabilities. It uses "
        "the Builder Pattern (ReportBuilder) to construct reports step-by-step and the Facade Pattern "
        "(DashboardFacade) to provide a simplified interface for aggregated dashboard data."
    )
    add_code_block(doc, read_file("backend/reports/builder.py"), "Source Code — reports/builder.py (Builder Pattern):")
    add_code_block(doc, read_file("backend/reports/services.py"), "Source Code — reports/services.py:")
    add_code_block(doc, read_file("backend/reports/dashboard.py"), "Source Code — reports/dashboard.py:")

    # 8.8 Core
    add_heading_styled(doc, "8.8 Core Module", level=2)
    add_body(doc,
        "The Core module contains shared utilities, design pattern implementations, and cross-cutting "
        "concerns used across the entire application. This includes the database connection, "
        "RBAC permission classes, and multiple design pattern implementations."
    )
    add_code_block(doc, read_file("backend/core/db.py"), "Source Code — core/db.py (Database Connection):")
    add_code_block(doc, read_file("backend/core/permissions.py"), "Source Code — core/permissions.py (RBAC Permissions):")
    add_code_block(doc, read_file("backend/core/services.py"), "Source Code — core/services.py (Facade Pattern):")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 9 – DESIGN PATTERNS
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "9. Design Patterns Implementation", level=1)
    add_body(doc,
        "DevFlow implements nine (9) industry-standard software design patterns categorized into "
        "Creational, Behavioral, and Structural types. These patterns promote code reusability, "
        "separation of concerns, and maintainability."
    )

    make_table(doc,
        ["#", "Pattern Name", "Category", "Implementation File", "Purpose in DevFlow"],
        [
            ["1", "Singleton", "Creational", "core/singleton.py", "Ensures single MongoDB connection instance"],
            ["2", "Factory", "Creational", "core/factories.py", "Creates User objects with specific roles"],
            ["3", "Builder", "Creational", "reports/builder.py", "Constructs reports step-by-step"],
            ["4", "Strategy", "Behavioral", "tasks/strategy.py", "Encapsulates task priority algorithms"],
            ["5", "State", "Behavioral", "tasks/state.py", "Manages task status transitions"],
            ["6", "Observer", "Behavioral", "core/observers.py, notifications/", "Event-driven notification system"],
            ["7", "Command", "Behavioral", "core/commands.py", "Encapsulates task operations as objects"],
            ["8", "Chain of Responsibility", "Behavioral", "core/approval_chain.py", "Approval workflow (PM → Admin)"],
            ["9", "Facade", "Structural", "core/services.py", "Simplified dashboard data aggregation"],
        ],
        col_widths=[1, 3.5, 2.5, 4.5, 5]
    )

    # 9.1 Singleton
    add_heading_styled(doc, "9.1 Singleton Pattern", level=2)
    add_body(doc,
        "The Singleton pattern ensures that a class has only one instance and provides a global point of "
        "access to it. In DevFlow, the MongoConnection class in core/singleton.py uses this pattern to "
        "prevent creating multiple database connections to MongoDB Atlas, which would waste resources."
    )
    add_body(doc, "How it works: The class maintains a static _instance variable. The get_instance() static method checks if _instance is None; if so, it creates a new MongoConnection object and stores it. Subsequent calls return the same instance.")
    add_code_block(doc, read_file("backend/core/singleton.py"), "Implementation — core/singleton.py:")

    # 9.2 Factory
    add_heading_styled(doc, "9.2 Factory Pattern", level=2)
    add_body(doc,
        "The Factory pattern provides an interface for creating objects without specifying the exact class "
        "of object that will be created. In DevFlow, UserFactory in core/factories.py creates User documents "
        "with the specified role, abstracting away the creation details."
    )
    add_code_block(doc, read_file("backend/core/factories.py"), "Implementation — core/factories.py:")

    # 9.3 Builder
    add_heading_styled(doc, "9.3 Builder Pattern", level=2)
    add_body(doc,
        "The Builder pattern separates the construction of a complex object from its representation. "
        "In DevFlow, the ReportBuilder class constructs a Report object step-by-step by adding project count, "
        "sprint count, task count, and bug count sequentially, then returns the final built report."
    )
    add_code_block(doc, read_file("backend/reports/builder.py"), "Implementation — reports/builder.py:")
    add_code_block(doc, read_file("backend/reports/services.py"), "Usage — reports/services.py:")

    # 9.4 Strategy
    add_heading_styled(doc, "9.4 Strategy Pattern", level=2)
    add_body(doc,
        "The Strategy pattern defines a family of algorithms, encapsulates each one, and makes them "
        "interchangeable. In DevFlow, three priority strategies (LowPriorityStrategy, MediumPriorityStrategy, "
        "HighPriorityStrategy) are encapsulated and a PriorityContext class dynamically selects which "
        "strategy to execute at runtime."
    )
    add_code_block(doc, read_file("backend/tasks/strategy.py"), "Implementation — tasks/strategy.py:")

    # 9.5 State
    add_heading_styled(doc, "9.5 State Pattern", level=2)
    add_body(doc,
        "The State pattern allows an object to alter its behavior when its internal state changes. "
        "In DevFlow, task status transitions are governed by state objects:\n"
        "• TodoState.next() → 'InProgress'\n"
        "• InProgressState.next() → 'Testing'\n"
        "• TestingState.next() → 'Done'\n"
        "• DoneState.next() → 'Done' (terminal state)\n\n"
        "This prevents invalid transitions and keeps the state logic clean and extensible."
    )
    add_code_block(doc, read_file("backend/tasks/state.py"), "Implementation — tasks/state.py:")

    # 9.6 Observer
    add_heading_styled(doc, "9.6 Observer Pattern", level=2)
    add_body(doc,
        "The Observer pattern defines a one-to-many dependency between objects so that when one object "
        "changes state, all its dependents are notified automatically. DevFlow has two Observer implementations:"
    )
    add_body(doc,
        "1. Core Observers (core/observers.py): Subject manages a list of observers (EmailObserver, "
        "NotificationObserver) and notifies all when an event occurs.\n"
        "2. Notification Observers (notifications/subject.py + observer.py): Domain-specific implementation "
        "for the notification subsystem."
    )
    add_code_block(doc, read_file("backend/core/observers.py"), "Implementation — core/observers.py:")
    add_code_block(doc, read_file("backend/notifications/subject.py"), "Implementation — notifications/subject.py:")
    add_code_block(doc, read_file("backend/notifications/observer.py"), "Implementation — notifications/observer.py:")

    # 9.7 Command
    add_heading_styled(doc, "9.7 Command Pattern", level=2)
    add_body(doc,
        "The Command pattern encapsulates a request as an object, allowing parameterization of clients "
        "with different requests, queuing, and logging. DevFlow defines a base Command class with "
        "CreateTaskCommand and DeleteTaskCommand as concrete implementations."
    )
    add_code_block(doc, read_file("backend/core/commands.py"), "Implementation — core/commands.py:")

    # 9.8 Chain of Responsibility
    add_heading_styled(doc, "9.8 Chain of Responsibility Pattern", level=2)
    add_body(doc,
        "The Chain of Responsibility pattern passes a request along a chain of handlers. Each handler "
        "decides either to process the request or to pass it to the next handler in the chain. "
        "In DevFlow, project approval requests pass through PMApproval first, and if not handled, "
        "escalate to AdminApproval."
    )
    add_code_block(doc, read_file("backend/core/approval_chain.py"), "Implementation — core/approval_chain.py:")

    # 9.9 Facade
    add_heading_styled(doc, "9.9 Facade Pattern", level=2)
    add_body(doc,
        "The Facade pattern provides a simplified interface to a complex subsystem. In DevFlow, "
        "DashboardFacade in core/services.py provides a single get_dashboard() method that internally "
        "queries four different collections (Projects, Sprints, Tasks, Bugs) and returns an aggregated "
        "result — hiding the complexity from the calling code."
    )
    add_code_block(doc, read_file("backend/core/services.py"), "Implementation — core/services.py:")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 10 – API DOCUMENTATION
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "10. REST API Documentation", level=1)
    add_body(doc,
        "DevFlow exposes a comprehensive RESTful API organized by module. All endpoints are prefixed "
        "with /api/ and return JSON responses. Authentication is required for most endpoints via JWT "
        "Bearer tokens in the Authorization header."
    )

    add_heading_styled(doc, "10.1 Accounts API", level=2)
    make_table(doc,
        ["Method", "Endpoint", "Auth", "Permission", "Description"],
        [
            ["POST", "/api/accounts/register/", "No", "Public", "Register a new user with name, email, password, role"],
            ["POST", "/api/accounts/login/", "No", "Public", "Authenticate and receive JWT tokens (access + refresh)"],
        ],
        col_widths=[2, 5, 1.5, 3, 5]
    )

    add_heading_styled(doc, "10.2 Projects API", level=2)
    make_table(doc,
        ["Method", "Endpoint", "Auth", "Permission", "Description"],
        [
            ["GET", "/api/projects/", "Yes", "Any authenticated", "List all projects"],
            ["POST", "/api/projects/create/", "Yes", "ProjectManager", "Create a new project"],
            ["GET", "/api/projects/<id>/", "Yes", "Any authenticated", "Get project details by ID"],
            ["PUT", "/api/projects/update/<id>/", "Yes", "Any authenticated", "Update project details"],
            ["DELETE", "/api/projects/delete/<id>/", "Yes", "Any authenticated", "Delete a project"],
        ],
        col_widths=[2, 5, 1.5, 3, 5]
    )

    add_heading_styled(doc, "10.3 Sprints API", level=2)
    make_table(doc,
        ["Method", "Endpoint", "Auth", "Permission", "Description"],
        [
            ["GET", "/api/sprints/", "Yes", "Any authenticated", "List all sprints"],
            ["POST", "/api/sprints/create/", "Yes", "Any authenticated", "Create a new sprint"],
            ["GET", "/api/sprints/<id>/", "Yes", "Any authenticated", "Get sprint details by ID"],
            ["PUT", "/api/sprints/update/<id>/", "Yes", "Any authenticated", "Update sprint details"],
            ["DELETE", "/api/sprints/delete/<id>/", "Yes", "Any authenticated", "Delete a sprint"],
            ["PUT", "/api/sprints/close/<id>/", "Yes", "Any authenticated", "Close/complete a sprint"],
        ],
        col_widths=[2, 5, 1.5, 3, 5]
    )

    add_heading_styled(doc, "10.4 Tasks API", level=2)
    make_table(doc,
        ["Method", "Endpoint", "Auth", "Permission", "Description"],
        [
            ["GET", "/api/tasks/", "Yes", "Any authenticated", "List all tasks"],
            ["POST", "/api/tasks/create/", "Yes", "Any authenticated", "Create a new task"],
            ["GET", "/api/tasks/<id>/", "Yes", "Any authenticated", "Get task details by ID"],
            ["PUT", "/api/tasks/update/<id>/", "Yes", "Any authenticated", "Update task details"],
            ["DELETE", "/api/tasks/delete/<id>/", "Yes", "Any authenticated", "Delete a task"],
            ["PUT", "/api/tasks/status/<id>/", "Yes", "Developer only", "Advance task to next status (State Pattern)"],
        ],
        col_widths=[2, 5, 1.5, 3, 5]
    )

    add_heading_styled(doc, "10.5 Bugs API", level=2)
    make_table(doc,
        ["Method", "Endpoint", "Auth", "Permission", "Description"],
        [
            ["GET", "/api/bugs/", "Yes", "Any authenticated", "List all bugs"],
            ["POST", "/api/bugs/create/", "Yes", "Any authenticated", "Report a new bug"],
            ["GET", "/api/bugs/<id>/", "Yes", "Any authenticated", "Get bug details by ID"],
            ["PUT", "/api/bugs/assign/<id>/", "Yes", "Any authenticated", "Assign a bug to a developer"],
            ["PUT", "/api/bugs/resolve/<id>/", "Yes", "Any authenticated", "Mark bug as resolved"],
            ["PUT", "/api/bugs/verify/<id>/", "Yes", "QA only", "Verify a resolved bug"],
            ["DELETE", "/api/bugs/delete/<id>/", "Yes", "Any authenticated", "Delete a bug"],
        ],
        col_widths=[2, 5, 1.5, 3, 5]
    )

    add_heading_styled(doc, "10.6 Reports & Dashboard API", level=2)
    make_table(doc,
        ["Method", "Endpoint", "Auth", "Permission", "Description"],
        [
            ["GET", "/api/reports/generate/", "Yes", "Any authenticated", "Generate project report (Builder Pattern)"],
            ["GET", "/api/reports/dashboard/", "Yes", "Any authenticated", "Get dashboard statistics (Facade Pattern)"],
        ],
        col_widths=[2, 5, 1.5, 3, 5]
    )

    add_heading_styled(doc, "10.7 Notifications API", level=2)
    make_table(doc,
        ["Method", "Endpoint", "Auth", "Permission", "Description"],
        [
            ["GET", "/api/notifications/", "Yes", "Any authenticated", "List all notifications"],
            ["PUT", "/api/notifications/<id>/read/", "Yes", "Any authenticated", "Mark a notification as read"],
        ],
        col_widths=[2, 5, 1.5, 3, 5]
    )

    add_heading_styled(doc, "10.8 Main URL Configuration", level=2)
    add_code_block(doc, read_file("backend/devflow/urls.py"), "Source Code — devflow/urls.py:")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 11 – AUTHENTICATION
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "11. Authentication & Authorization (JWT)", level=1)
    add_body(doc,
        "DevFlow uses JSON Web Tokens (JWT) for stateless authentication. The system generates both "
        "access and refresh tokens upon successful login. The access token contains the user_id and "
        "role claims, which are used by the custom MongoJWTAuthentication backend to authenticate "
        "API requests."
    )

    add_heading_styled(doc, "11.1 Authentication Flow", level=2)
    auth_flow = [
        "1. User submits email and password to POST /api/accounts/login/",
        "2. Backend verifies credentials using bcrypt password comparison",
        "3. On success, JWTService creates access + refresh tokens with user_id and role claims",
        "4. Tokens are returned as JSON response to the frontend",
        "5. Frontend stores the access token in localStorage and as an HTTP cookie",
        "6. For subsequent API calls, the frontend includes the token in the Authorization header: Bearer <token>",
        "7. The backend's MongoJWTAuthentication class extracts and validates the token",
        "8. The user_id claim is used to look up the MongoDB user document",
        "9. The authenticated user object is attached to request.user for permission checks",
    ]
    for step in auth_flow:
        add_bullet(doc, step)

    add_heading_styled(doc, "11.2 Token Structure", level=2)
    make_table(doc,
        ["Token Type", "Contains", "Lifetime", "Purpose"],
        [
            ["Access Token", "user_id, role, exp, iat", "Short-lived (configurable)", "API authentication for protected endpoints"],
            ["Refresh Token", "user_id, role, exp, iat", "Long-lived (configurable)", "Used to obtain new access tokens"],
        ],
        col_widths=[3.5, 4.5, 4, 4.5]
    )

    add_heading_styled(doc, "11.3 Password Security", level=2)
    add_body(doc,
        "Passwords are hashed using bcrypt before storage. The PasswordService class provides two static methods:\n"
        "• hash_password(password): Generates a bcrypt hash with automatic salt generation.\n"
        "• verify_password(password, hashed): Compares a plaintext password against the stored hash.\n\n"
        "Bcrypt is a computationally expensive algorithm specifically designed for password hashing, "
        "providing protection against brute-force and rainbow table attacks."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 12 – RBAC
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "12. Role-Based Access Control (RBAC)", level=1)
    add_body(doc,
        "DevFlow implements a comprehensive Role-Based Access Control (RBAC) system with five distinct "
        "user roles. Each role has specific permissions that determine what features and actions are "
        "accessible. Permissions are enforced at the API level using custom DRF permission classes "
        "that inherit from BasePermission."
    )

    add_heading_styled(doc, "12.1 User Roles", level=2)
    make_table(doc,
        ["Role", "Description", "Key Permissions"],
        [
            ["Admin", "System administrator with full control", "Manage users, view reports, full system access"],
            ["ProjectManager", "Manages projects and sprint planning", "Create projects, create sprints, create tasks, view reports"],
            ["Developer", "Development team member", "Update task status (Kanban workflow), create bugs"],
            ["QA", "Quality Assurance engineer", "Verify bugs, update task status"],
            ["Client", "External stakeholder with view access", "View reports (read-only access)"],
        ],
        col_widths=[3.5, 5, 8]
    )

    add_heading_styled(doc, "12.2 Permission Matrix", level=2)
    add_body(doc, "The following matrix shows which features are accessible to each role (✔ = allowed, ✘ = denied):")
    make_table(doc,
        ["Feature", "Admin", "PM", "Dev", "QA", "Client"],
        [
            ["Manage Users", "✔", "✘", "✘", "✘", "✘"],
            ["Create Project", "✘", "✔", "✘", "✘", "✘"],
            ["Create Sprint", "✘", "✔", "✔", "✘", "✘"],
            ["Create Task", "✘", "✔", "✔", "✘", "✘"],
            ["Update Task Status", "✘", "✘", "✔", "✔", "✘"],
            ["Create Bug", "✘", "✘", "✔", "✔", "✘"],
            ["Verify Bug", "✘", "✘", "✘", "✔", "✘"],
            ["View Reports", "✔", "✔", "✘", "✘", "✔"],
        ],
        col_widths=[4, 2.5, 2.5, 2.5, 2.5, 2.5]
    )

    add_heading_styled(doc, "12.3 Permission Implementation", level=2)
    add_body(doc,
        "Permissions are enforced using custom DRF permission classes that check the request.user.role "
        "attribute. Each class returns True only if the user has the required role."
    )
    add_code_block(doc, read_file("backend/core/permissions.py"), "Source Code — core/permissions.py:")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 13 – FRONTEND ARCHITECTURE
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "13. Frontend Architecture (Next.js)", level=1)
    add_body(doc,
        "The DevFlow frontend is built using Next.js 16 with the App Router, React 19, TypeScript, "
        "and TailwindCSS v4. The frontend communicates with the Django backend via RESTful API calls "
        "using the native fetch API with JWT Bearer token authentication."
    )

    # 13.1
    add_heading_styled(doc, "13.1 Pages & Routing", level=2)
    make_table(doc,
        ["Route", "File", "Authentication", "Description"],
        [
            ["/", "app/page.tsx", "Required", "Home page (auto-redirects to /dashboard)"],
            ["/login", "app/login/page.tsx", "Public", "User login form with JWT authentication"],
            ["/register", "app/register/page.tsx", "Public", "User registration form with role selection"],
            ["/dashboard", "app/dashboard/page.tsx", "Required", "Main dashboard with stats, charts, and quick actions"],
            ["/projects", "app/projects/page.tsx", "Required", "Project management (CRUD operations)"],
            ["/tasks", "app/tasks/page.tsx", "Required", "Task management with Kanban workflow"],
            ["/bugs", "app/bugs/page.tsx", "Required", "Bug tracking and lifecycle management"],
        ],
        col_widths=[3, 4.5, 3, 6]
    )

    # 13.2
    add_heading_styled(doc, "13.2 Components", level=2)
    make_table(doc,
        ["Component", "File", "Description"],
        [
            ["Navbar", "components/Navbar.tsx", "Top navigation bar with DevFlow logo, status indicator, and logout button"],
            ["Sidebar", "components/Sidebar.tsx", "Left sidebar with navigation menu (Dashboard, Projects, Tasks, Bugs) and user info"],
            ["Icons", "components/Icons.tsx", "11 custom SVG icon components (Dashboard, Folder, Tasks, Bug, Trash, Plus, etc.)"],
            ["ProjectCard", "components/ProjectCard.tsx", "Card component for displaying project information"],
            ["TaskCard", "components/TaskCard.tsx", "Card component for displaying task information"],
            ["BugCard", "components/BugCard.tsx", "Card component for displaying bug information"],
        ],
        col_widths=[3, 4.5, 9]
    )

    # 13.3
    add_heading_styled(doc, "13.3 Services Layer", level=2)
    add_body(doc,
        "The services layer provides abstracted API communication functions. Each service file corresponds "
        "to a backend module and encapsulates HTTP calls with proper authentication headers."
    )
    make_table(doc,
        ["Service File", "Functions", "Backend Module"],
        [
            ["api.ts", "API_URL constant (http://localhost:8000/api)", "Base configuration"],
            ["auth.ts", "login(), register(), getAuthHeaders()", "accounts/"],
            ["project.ts", "getProjects(), createProject(), updateProject(), deleteProject()", "projects/"],
            ["task.ts", "getTasks(), createTask(), changeTaskStatus(), updateTask(), deleteTask()", "tasks/"],
            ["bug.ts", "getBugs(), createBug(), assignBug(), resolveBug(), verifyBug(), deleteBug()", "bugs/"],
        ],
        col_widths=[3, 7, 6.5]
    )

    # 13.4
    add_heading_styled(doc, "13.4 Middleware & Auth Guard", level=2)
    add_body(doc,
        "The frontend implements a Next.js middleware (middleware.ts) that acts as an authentication guard. "
        "It intercepts all requests and checks for a JWT token in the cookies. If the token is missing "
        "and the route is not public (/login, /register, /_next, /api), the user is redirected to the "
        "login page. This ensures that only authenticated users can access protected pages."
    )
    add_code_block(doc, read_file("frontend/middleware.ts"), "Source Code — middleware.ts:")

    # 13.5
    add_heading_styled(doc, "13.5 UI/UX Design System", level=2)
    add_body(doc,
        "DevFlow features a premium dark-themed UI with the following design characteristics:"
    )
    add_bullet(doc, "Dark Background: Deep navy gradient (#060814 → #0a0e29) for a modern, immersive feel.", "Color Palette — ")
    add_bullet(doc, "Glassmorphism Panels: Semi-transparent cards with backdrop blur effects.", "Glass Panels — ")
    add_bullet(doc, "Neon Accent Colors: Indigo (#818cf8), Teal (#2dd4bf), Emerald (#34d399), Amber (#fbbf24), Rose (#f87171).", "Neon Theme — ")
    add_bullet(doc, "Plus Jakarta Sans (body text) and JetBrains Mono (code) for premium readability.", "Typography — ")
    add_bullet(doc, "Floating gradient orbs, pulse animations, hover transitions, and smooth micro-interactions.", "Animations — ")
    add_bullet(doc, "Custom styled scrollbars, autofill handling for dark fields, and responsive layouts.", "Polished Details — ")
    add_bullet(doc, "SVG Donut chart for resource distribution and Bezier area chart for delivery velocity.", "Dashboard Charts — ")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 14 – DOCKER
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "14. Containerization & Deployment (Docker)", level=1)
    add_body(doc,
        "DevFlow is containerized using Docker for consistent deployment across environments. "
        "The application uses a Dockerfile for the Python backend and docker-compose.yml to "
        "orchestrate multi-service deployment with MongoDB."
    )

    add_heading_styled(doc, "14.1 Dockerfile", level=2)
    add_body(doc,
        "The Dockerfile builds a Python 3.12 image, copies the application code, installs dependencies "
        "from requirements.txt, exposes port 8000, and starts the Django development server."
    )
    add_code_block(doc, read_file("Dockerfile"), "Dockerfile:")

    add_heading_styled(doc, "14.2 Docker Compose", level=2)
    add_body(doc,
        "Docker Compose orchestrates two services:\n"
        "1. backend: The Django application (built from Dockerfile, exposed on port 8000)\n"
        "2. mongodb: Official MongoDB image (exposed on port 27017)"
    )
    add_code_block(doc, read_file("docker-compose.yml"), "docker-compose.yml:")

    add_heading_styled(doc, "14.3 Deployment Commands", level=2)
    deploy_cmds = """# Build and start all services
docker-compose up --build

# Start in detached (background) mode
docker-compose up -d

# Stop all services
docker-compose down

# View running containers
docker ps

# View logs
docker-compose logs -f backend"""
    add_code_block(doc, deploy_cmds, "Docker Commands:")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 15 – ENVIRONMENT CONFIG
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "15. Environment Configuration", level=1)
    add_body(doc,
        "DevFlow uses environment variables managed through a .env file for sensitive configuration. "
        "The python-dotenv library loads these variables at application startup."
    )

    make_table(doc,
        ["Variable", "Purpose", "Example Value"],
        [
            ["MONGO_URI", "MongoDB Atlas connection string with TLS & replica set", "mongodb://user:pass@host:port/devflow?..."],
            ["SECRET_KEY", "Django secret key for cryptographic signing", "your_secret_key"],
            ["DEBUG", "Django debug mode flag", "True"],
        ],
        col_widths=[3.5, 6, 7]
    )

    add_heading_styled(doc, "15.1 Django Settings Configuration", level=2)
    add_code_block(doc, read_file("backend/devflow/settings.py"), "Source Code — devflow/settings.py:")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 16 – TESTING
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "16. Testing Strategy", level=1)
    add_body(doc, "DevFlow employs the following testing approaches to ensure system reliability:")

    add_heading_styled(doc, "16.1 API Testing", level=2)
    add_body(doc,
        "All REST API endpoints are tested using Postman or similar API testing tools. Each endpoint "
        "is verified for:\n"
        "• Correct HTTP status codes (200, 201, 400, 401, 403, 404)\n"
        "• Proper JSON response structure\n"
        "• Authentication enforcement (401 for missing tokens)\n"
        "• Permission enforcement (403 for unauthorized roles)\n"
        "• Data persistence verification in MongoDB"
    )

    add_heading_styled(doc, "16.2 Database Connection Testing", level=2)
    add_code_block(doc, read_file("backend/test_db.py"), "Source Code — test_db.py:")

    add_heading_styled(doc, "16.3 Frontend Testing", level=2)
    add_body(doc,
        "Frontend testing includes:\n"
        "• Manual UI testing across different browsers and screen sizes\n"
        "• Authentication flow testing (login, register, logout, token expiry)\n"
        "• CRUD operation verification for projects, tasks, and bugs\n"
        "• Role-based UI element visibility testing\n"
        "• Responsive design testing on various viewport sizes"
    )

    add_heading_styled(doc, "16.4 Design Pattern Verification", level=2)
    add_body(doc,
        "Each design pattern implementation is verified through:\n"
        "• Singleton: Confirmed single instance behavior via multiple get_instance() calls\n"
        "• State: Verified correct state transitions (Todo → InProgress → Testing → Done)\n"
        "• Strategy: Validated priority strategy selection and execution\n"
        "• Observer: Confirmed notification propagation to all attached observers\n"
        "• Builder: Verified step-by-step report construction and final output\n"
        "• Chain of Responsibility: Tested approval chain escalation from PM to Admin\n"
        "• Factory: Confirmed user creation with correct roles\n"
        "• Command: Verified task operation encapsulation and execution\n"
        "• Facade: Validated aggregated dashboard data retrieval"
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 17 – CONCLUSION
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "17. Conclusion", level=1)
    add_body(doc,
        "DevFlow is a comprehensive, production-grade Agile Project Management System that demonstrates "
        "mastery of modern full-stack web development principles, software design patterns, and enterprise "
        "architecture practices. The system successfully integrates a Django REST Framework backend with "
        "MongoDB Atlas and a Next.js 16 frontend to deliver a complete project management solution."
    )
    add_body(doc,
        "The implementation of nine (9) design patterns — Singleton, Factory, Builder, Strategy, State, "
        "Observer, Command, Chain of Responsibility, and Facade — demonstrates a strong understanding of "
        "object-oriented design principles and their practical application in real-world software systems. "
        "Each pattern serves a specific architectural purpose, from ensuring single database connections "
        "(Singleton) to managing complex task state workflows (State Pattern)."
    )
    add_body(doc,
        "The Role-Based Access Control (RBAC) system with five distinct roles (Admin, Project Manager, "
        "Developer, QA, Client) provides fine-grained security enforcement at both the API and UI layers. "
        "JWT-based authentication ensures secure, stateless communication between the frontend and backend."
    )
    add_body(doc,
        "The glassmorphism-themed frontend delivers a premium, modern user experience with smooth "
        "animations, interactive charts, and responsive design. Docker containerization enables consistent "
        "deployment across environments. Overall, DevFlow represents a well-architected, feature-complete "
        "software system that follows industry best practices and is ready for further enhancement."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 18 – FUTURE ENHANCEMENTS
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "18. Future Enhancements", level=1)
    add_body(doc, "The following enhancements are planned for future iterations of DevFlow:")

    enhancements = [
        ("Real-time Notifications: ", "Implement WebSocket-based real-time notifications using Django Channels to replace the current print-based Observer implementation."),
        ("Advanced Analytics: ", "Add time-series charts with actual historical data, burndown charts, velocity tracking, and team performance metrics."),
        ("File Attachments: ", "Support file uploads (documents, screenshots) for tasks and bugs using cloud storage integration (AWS S3 or similar)."),
        ("CI/CD Pipeline: ", "Integrate GitHub Actions or Jenkins for automated testing, building, and deployment."),
        ("Mobile Application: ", "Develop React Native mobile applications for iOS and Android platforms."),
        ("Multi-Tenant Support: ", "Enable organization-level isolation for enterprise SaaS deployment."),
        ("Audit Logging: ", "Implement comprehensive audit trails for all CRUD operations with user, timestamp, and action tracking."),
        ("Email Integration: ", "Complete the email notification service with SMTP configuration for real email delivery."),
        ("Search & Filtering: ", "Add advanced search and filtering capabilities across projects, tasks, and bugs."),
        ("Data Export: ", "Enable export of reports and data in PDF, CSV, and Excel formats."),
    ]
    for prefix, text in enhancements:
        add_bullet(doc, text, prefix)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 19 – REFERENCES
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "19. References", level=1)
    refs = [
        "Django Documentation — https://docs.djangoproject.com/",
        "Django REST Framework — https://www.django-rest-framework.org/",
        "MongoEngine Documentation — https://docs.mongoengine.org/",
        "MongoDB Atlas Documentation — https://www.mongodb.com/docs/atlas/",
        "Next.js Documentation — https://nextjs.org/docs",
        "React Documentation — https://react.dev/",
        "TailwindCSS Documentation — https://tailwindcss.com/docs",
        "Simple JWT (djangorestframework-simplejwt) — https://django-rest-framework-simplejwt.readthedocs.io/",
        "Docker Documentation — https://docs.docker.com/",
        "Bcrypt — https://pypi.org/project/bcrypt/",
        "Design Patterns: Elements of Reusable Object-Oriented Software — Gamma, Helm, Johnson, Vlissides (Gang of Four)",
        "Python-dotenv — https://pypi.org/project/python-dotenv/",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"[{i}]  {ref}")
        p.paragraph_format.space_after = Pt(4)

    # ── Save ──
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "DevFlow_Project_Report.docx")
    doc.save(output_path)
    print(f"\nReport generated successfully: {output_path}")
    print(f"Total chapters: 19")
    print(f"Design patterns documented: 9")
    print(f"API endpoints documented: 28")
    print(f"Database collections documented: 6")


if __name__ == "__main__":
    build_report()
